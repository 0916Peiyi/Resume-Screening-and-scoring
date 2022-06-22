import pdfplumber
from docx import Document
import re
import nltk
import os
nltk.download('punkt')
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords


def pdf_to_text(path):
    with pdfplumber.open(path) as pdf:
        content = ''
        for page in pdf.pages:
            content += ''.join(page.extract_text().split('\n')[:-1])
    return content


def docx_to_text(path):
    document = Document(path)
    # 有的简历是表格式样的，因此，不仅需要提取正文，还要提取表格
    col_keys = [] # 获取列名
    col_values = [] # 获取列值
    index_num = 0
    # 表格提取中，需要添加一个去重机制
    fore_str = ""
    cell_text = ""
    for table in document.tables:
        for row_index,row in enumerate(table.rows):
            for col_index,cell in enumerate(row.cells):
                if fore_str != cell.text:
                    if index_num % 2==0:
                        col_keys.append(cell.text)
                    else:
                        col_values.append(cell.text)
                    fore_str = cell.text
                    index_num +=1
                    cell_text += cell.text + ' '
    # 提取正文文本
    paragraphs_text = ""
    for paragraph in document.paragraphs:
        # 拼接一个list,包括段落的结构和内容
        paragraphs_text += paragraph.text + " "
    return cell_text, paragraphs_text


def read_file(path):
    file_type = os.path.splitext(path)[-1]
    if file_type[0] != '.':
        ValueError('Please insert a file with .docx or .pdf')
    file_type = file_type[1:]
    print(file_type)
    if file_type == 'pdf':
        return pdf_to_text(path)
    elif file_type == 'docx':
        cell_text, paragraphs_text = docx_to_text(path)
        return cell_text + paragraphs_text
    else:ValueError('Please insert a file with .docx or .pdf')

def clean_data(x, additional_chars):
    x = re.sub('{IMG:.?.?.?}', ' ', x)
    x = re.sub('<!--IMG_\d+-->', ' ', x)
    x = re.sub('(https?|ftp|file)://[-A-Za-z0-9+&@#/%?=~_|!:,.;]+[-A-Za-z0-9+&@#/%=~_|]', ' ', x)  # 过滤网址
    x = re.sub('<a[^>]*>', '', x).replace("</a>", " ")  # 过滤a标签
    x = re.sub('<P[^>]*>', '', x).replace("</P>", " ")  # 过滤P标签
    x = re.sub('<strong[^>]*>', ',', x).replace("</strong>", " ")  # 过滤strong标签
    x = re.sub('<br>', ',', x)  # 过滤br标签
    x = re.sub('www.[-A-Za-z0-9+&@#/%?=~_|!:,.;]+[-A-Za-z0-9+&@#/%=~_|]', '', x).replace("()", " ")  # 过滤www开头的网址
    x = re.sub('\s', ' ', x)   # 过滤不可见字符
    x = re.sub('Ⅴ', 'V', x)

    for wbad in additional_chars:
        x = x.replace(wbad, ' ')
    return x

    return x

def extract_tokenize_from_doc(path):
    content = read_file(path)
    additional_chars = set()
    additional_chars.update(re.findall(u'[^\u4e00-\u9fa5a-zA-Z0-9\*]', str(content)))
    # 一些需要保留的符号
    extra_chars = set("!#+._#‘’")
    additional_chars = additional_chars.difference(extra_chars)
    content = content.lower()
    content = clean_data(content, additional_chars)
    tokenizes = word_tokenize(content)
    # remove stop words
    tokenizes = [w for w in tokenizes if w not in stopwords.words('english')]

    return tokenizes



